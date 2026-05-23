"""Workflow Report Generator — generates DOCX/PDF/ODF reports for debate sessions.

Produces a structured report including:
- Debate title, case description, metadata
- Per-round agent outputs (full transcript)
- Consensus progression
- Final consensus and result
- Audit trail table (optional, if entries exist)
"""

from __future__ import annotations

import ast
import asyncio
import html as html_mod
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from weasyprint import HTML

from backend.workflow.audit_logger import AuditLogger
from backend.workflow.state_snapshot import StateSnapshotStore

logger = logging.getLogger(__name__)

_REPORTS_DIR = Path("reports")


def _display_agent_role(role: str) -> str:
    """Format an agent role for display.

    MVP debates pass pre-formatted names like ``"Strategist (deepseek-v4-flash)"``
    while legacy debates pass raw role names like ``"critic"``.  Capitalize
    only when the role starts with a lowercase letter.
    """
    if not role:
        return "Unbekannt"
    return role if role[0].isupper() else role.capitalize()


def _build_mvp_rounds_from_snapshot(debate_data: dict[str, Any]) -> list[dict[str, Any]]:
    """Build round data from state snapshot for MVP debates.

    MVP debates don't populate ``debate_data["rounds"]`` — the agent
    outputs live exclusively in the state snapshot's ``node_outputs``.
    This helper loads the snapshot, resolves agent names and LLM
    profiles from ``node_configs``, and returns a list of round dicts
    compatible with the existing renderer expectations.

    Returns:
        A list of round dicts with keys ``round``, ``consensus``,
        ``agent_outputs``.
    """
    session_id = debate_data.get("session_id", "")
    if not session_id:
        return []

    try:
        snap_store = StateSnapshotStore()
        snapshot = snap_store.get_latest(session_id)
    except Exception:
        logger.warning("Could not load state snapshot for session %s", session_id)
        return []

    if not snapshot:
        return []

    state = snapshot.get("state", {})
    node_outputs: list[dict] = state.get("node_outputs", [])
    node_configs_raw: dict[str, Any] = state.get("node_configs", {})
    llm_assignments: dict[str, str] = debate_data.get("llm_assignments", {})

    if not node_outputs:
        return []

    # Parse node_configs — values may be stored as Python repr strings
    config_by_node: dict[str, dict] = {}
    for nid, cfg in node_configs_raw.items():
        if isinstance(cfg, str):
            try:
                cfg = ast.literal_eval(cfg)
            except (ValueError, SyntaxError):
                cfg = {}
        config_by_node[nid] = cfg if isinstance(cfg, dict) else {}

    # Build agent outputs with proper names and metadata
    agent_outputs: list[dict[str, Any]] = []
    for no in node_outputs:
        role = no.get("role", "")
        node_id = no.get("node_id", "")
        config = config_by_node.get(node_id, {})

        llm_model = config.get("llm_model", "")
        llm_profile_id = config.get("llm_profile_id", "") or llm_assignments.get(role, "")
        role_type_name = config.get("role_type_name", "")

        # Build a human-readable agent name: "Strategist (deepseek-v4-flash)"
        agent_name = role_type_name or role.replace("_", " ").title()
        if llm_model:
            agent_name = f"{agent_name} ({llm_model})"

        agent_outputs.append(
            {
                "role": agent_name,
                "content": no.get("content", ""),
                "tokens_used": no.get("tokens_used", 0),
                "duration_ms": no.get("duration_ms", 0),
                "llm_profile_id": llm_profile_id,
                "round": no.get("round"),
            }
        )

    current_round = state.get("current_round", debate_data.get("current_round", 1))

    return [
        {
            "round": current_round,
            "consensus": state.get("final_consensus", debate_data.get("result", {}).get("consensus", 0.0)),
            "agent_outputs": agent_outputs,
        }
    ]


class WorkflowReportGenerator:
    """Generates structured reports for completed debate sessions."""

    def __init__(self, db_path: Path | str | None = None) -> None:
        self._db_path = Path(db_path) if db_path else None
        self._audit = AuditLogger(self._db_path)
        _REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    async def generate(
        self,
        session_id: str,
        fmt: str = "docx",
        debate_data: dict[str, Any] | None = None,
    ) -> Path:
        """Generate a report for *session_id* in the given *fmt*.

        Args:
            session_id: The workflow session ID (same as debate_id).
            fmt: Output format — ``"docx"``, ``"pdf"``, or ``"odf"``.
            debate_data: Optional debate dict from DebateStore.  When
                provided the report includes the full debate transcript.

        Returns:
            Path to the generated report file.
        """
        if fmt not in ("docx", "pdf", "odf"):
            raise ValueError(f"Unsupported format: {fmt!r}")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{session_id[:8]}_{ts}.{fmt}"
        path = _REPORTS_DIR / filename

        # Load audit data (supplementary)
        audit_entries = self._audit.get_audit_log_for_replay(session_id)

        # Build transcript data from debate store
        transcript = self._build_transcript(debate_data)

        if fmt == "docx":
            await asyncio.to_thread(self._build_docx, session_id, transcript, audit_entries, path)
        elif fmt == "pdf":
            await asyncio.to_thread(self._build_pdf, session_id, transcript, audit_entries, path)
        elif fmt == "odf":
            await asyncio.to_thread(self._build_odf, session_id, transcript, audit_entries, path)

        logger.info("Report generated: %s", path)
        return path

    # ------------------------------------------------------------------
    # Transcript extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _build_transcript(debate_data: dict[str, Any] | None) -> dict[str, Any]:
        """Extract structured transcript data from a debate dict.

        Returns a normalized dict with keys: title, case_text, language,
        model, max_rounds, threshold, status, current_round,
        final_consensus, rounds (list of round dicts).
        """
        if not debate_data:
            return {
                "title": "",
                "case_text": "",
                "language": "de",
                "model": "",
                "max_rounds": 0,
                "threshold": 0.0,
                "status": "unknown",
                "current_round": 0,
                "final_consensus": 0.0,
                "rounds": [],
                "output": "",
            }

        req = debate_data.get("request", {})
        result = debate_data.get("result", {})
        rounds = debate_data.get("rounds", [])

        # rounds may come from either the debate store or the result dict
        if not rounds and result:
            rounds = result.get("rounds", [])

        # MVP debates: build rounds from state snapshot node_outputs
        is_mvp = debate_data.get("is_mvp", False)
        if is_mvp and not rounds:
            rounds = _build_mvp_rounds_from_snapshot(debate_data)

        return {
            "title": debate_data.get("title", ""),
            "case_text": req.get("case", {}).get("text", "") if isinstance(req.get("case"), dict) else str(req.get("case", "")),
            "language": req.get("language", "de"),
            "model": req.get("llm_profile_id", ""),
            "max_rounds": req.get("max_rounds", 0),
            "threshold": req.get("consensus_threshold", 0.0),
            "status": str(debate_data.get("status", "unknown")),
            "current_round": debate_data.get("current_round", result.get("current_round", 0)),
            "final_consensus": result.get("final_consensus", debate_data.get("final_consensus", 0.0)),
            "rounds": rounds,
            "output": result.get("output", ""),
            "is_mvp": is_mvp,
        }

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _build_docx(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"

        # --- Title ---
        title = transcript["title"] or f"Debatte {session_id[:8]}"
        doc.add_heading(title, level=0)

        # --- Metadata ---
        doc.add_heading("Zusammenfassung", level=1)
        table_meta = doc.add_table(rows=0, cols=2, style="Table Grid")
        meta_rows = [
            ("Session-ID", session_id),
            ("Titel", transcript["title"]),
            ("Sprache", transcript["language"]),
            ("Modell", transcript["model"]),
            ("Max. Runden", str(transcript["max_rounds"])),
            ("Konsens-Schwelle", f"{transcript['threshold']:.0%}"),
            ("Status", transcript["status"]),
            ("Runden absolviert", str(transcript["current_round"])),
            ("Finaler Konsens", f"{transcript['final_consensus']:.1%}"),
            ("Generiert", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        for label, value in meta_rows:
            row = table_meta.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # --- Case description ---
        case_text = transcript["case_text"]
        if case_text:
            doc.add_heading("Fallbeschreibung", level=1)
            doc.add_paragraph(case_text)

        # --- Debate transcript (per round) ---
        rounds = transcript["rounds"]
        if rounds:
            doc.add_heading("Debatte-Transkript", level=1)
            for rd in rounds:
                round_num = rd.get("round", "?")
                consensus = rd.get("consensus", 0.0)
                doc.add_heading(f"Runde {round_num} — Konsens: {consensus:.1%}", level=2)

                agent_outputs = rd.get("agent_outputs", [])
                for ao in agent_outputs:
                    role = ao.get("role", "unbekannt")
                    content = ao.get("content", "")
                    tokens = ao.get("tokens_used", 0)
                    duration_ms = ao.get("duration_ms", 0)
                    llm_pid = ao.get("llm_profile_id", "")
                    heading = _display_agent_role(role)
                    if llm_pid:
                        heading += f" — {llm_pid}"
                    doc.add_heading(heading, level=3)
                    if content:
                        for paragraph in content.split("\n\n"):
                            paragraph = paragraph.strip()
                            if paragraph:
                                doc.add_paragraph(paragraph)
                    meta_parts = []
                    if tokens:
                        meta_parts.append(f"Tokens: {tokens}")
                    if duration_ms:
                        meta_parts.append(f"Latenz: {duration_ms}ms")
                    if llm_pid:
                        meta_parts.append(f"LLM-Profil: {llm_pid}")
                    if meta_parts:
                        doc.add_paragraph(" | ".join(meta_parts)).italic = True

        # --- Final output ---
        output = transcript["output"]
        if output:
            doc.add_heading("Ergebnis", level=1)
            for paragraph in output.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)

        # --- Audit trail (supplementary) ---
        if audit_entries:
            doc.add_heading("Audit-Trail", level=1)
            table = doc.add_table(rows=1, cols=7, style="Table Grid")
            headers = ["Zeitstempel", "Ereignis", "Knoten", "Aktor", "LLM-Profil", "Latenz (ms)", "Tokens"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for entry in audit_entries:
                row = table.add_row()
                row.cells[0].text = str(entry.get("timestamp", ""))
                row.cells[1].text = str(entry.get("event_type", ""))
                row.cells[2].text = str(entry.get("node_id", ""))
                row.cells[3].text = str(entry.get("actor", ""))
                row.cells[4].text = str(entry.get("llm_profile_id", ""))
                row.cells[5].text = str(entry.get("latency_ms", 0))
                row.cells[6].text = str(entry.get("prompt_tokens", 0) + entry.get("completion_tokens", 0))

        doc.save(str(path))

    # ------------------------------------------------------------------
    # PDF (via WeasyPrint)
    # ------------------------------------------------------------------

    def _build_pdf(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        html_content = self._render_html(session_id, transcript, audit_entries)
        HTML(string=html_content).write_pdf(str(path))

    # ------------------------------------------------------------------
    # ODF
    # ------------------------------------------------------------------

    def _build_odf(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        html_content = self._render_html(session_id, transcript, audit_entries)
        try:
            from odf.opendocument import OpenDocumentText
            from odf.text import H, P

            doc = OpenDocumentText()

            title = transcript["title"] or f"Debatte {session_id[:8]}"
            doc.text.addElement(H(text=title, outlinelevel=1))

            # Metadata
            doc.text.addElement(H(text="Zusammenfassung", outlinelevel=2))
            meta_lines = [
                f"Session-ID: {session_id}",
                f"Titel: {transcript['title']}",
                f"Sprache: {transcript['language']}",
                f"Modell: {transcript['model']}",
                f"Max. Runden: {transcript['max_rounds']}",
                f"Konsens-Schwelle: {transcript['threshold']:.0%}",
                f"Status: {transcript['status']}",
                f"Runden absolviert: {transcript['current_round']}",
                f"Finaler Konsens: {transcript['final_consensus']:.1%}",
            ]
            for line in meta_lines:
                doc.text.addElement(P(text=line))

            # Case description
            case_text = transcript["case_text"]
            if case_text:
                doc.text.addElement(H(text="Fallbeschreibung", outlinelevel=2))
                for para in case_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.text.addElement(P(text=para))

            # Transcript
            rounds = transcript["rounds"]
            if rounds:
                doc.text.addElement(H(text="Debatte-Transkript", outlinelevel=2))
                for rd in rounds:
                    round_num = rd.get("round", "?")
                    consensus = rd.get("consensus", 0.0)
                    doc.text.addElement(H(text=f"Runde {round_num} — Konsens: {consensus:.1%}", outlinelevel=3))
                    for ao in rd.get("agent_outputs", []):
                        role = ao.get("role", "unbekannt")
                        content = ao.get("content", "")
                        llm_pid = ao.get("llm_profile_id", "")
                        role_label = _display_agent_role(role)
                        if llm_pid:
                            role_label += f" — {llm_pid}"
                        doc.text.addElement(P(text=f"[{role_label}]"))
                        if content:
                            for para in content.split("\n\n"):
                                para = para.strip()
                                if para:
                                    doc.text.addElement(P(text=para))

            # Final output
            output = transcript["output"]
            if output:
                doc.text.addElement(H(text="Ergebnis", outlinelevel=2))
                for para in output.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.text.addElement(P(text=para))

            # Audit trail
            if audit_entries:
                doc.text.addElement(H(text="Audit-Trail", outlinelevel=2))
                for entry in audit_entries:
                    llm_pid = entry.get("llm_profile_id", "")
                    line = f"{entry.get('timestamp', '')} | {entry.get('event_type', '')} | {entry.get('node_id', '')} | {entry.get('actor', '')}"
                    if llm_pid:
                        line += f" | LLM: {llm_pid}"
                    doc.text.addElement(P(text=line))

            doc.save(str(path))
        except ImportError:
            # Fallback: write HTML if odfpy not available
            path = path.with_suffix(".html")
            path.write_text(html_content, encoding="utf-8")
            logger.warning("odfpy not available; wrote HTML instead: %s", path)

    # ------------------------------------------------------------------
    # HTML rendering (shared by PDF and ODF)
    # ------------------------------------------------------------------

    @staticmethod
    def _render_html(
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
    ) -> str:
        esc = html_mod.escape
        title = esc(transcript["title"] or f"Debatte {session_id[:8]}")

        # --- Metadata table ---
        meta_pairs = [
            ("Session-ID", session_id),
            ("Titel", esc(transcript["title"])),
            ("Sprache", esc(transcript["language"])),
            ("Modell", esc(transcript["model"])),
            ("Max. Runden", str(transcript["max_rounds"])),
            ("Konsens-Schwelle", f"{transcript['threshold']:.0%}"),
            ("Status", esc(transcript["status"])),
            ("Runden absolviert", str(transcript["current_round"])),
            ("Finaler Konsens", f"{transcript['final_consensus']:.1%}"),
            ("Generiert", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        meta_rows = "".join(f"<tr><td><strong>{k}</strong></td><td>{v}</td></tr>" for k, v in meta_pairs)

        # --- Case description ---
        case_text = transcript["case_text"]
        case_section = ""
        if case_text:
            case_section = f'<h2>Fallbeschreibung</h2><div class="case-text">{esc(case_text)}</div>'

        # --- Rounds ---
        rounds_html = ""
        rounds = transcript["rounds"]
        if rounds:
            rounds_html += "<h2>Debatte-Transkript</h2>"
            for rd in rounds:
                round_num = rd.get("round", "?")
                consensus = rd.get("consensus", 0.0)
                rounds_html += f"<h3>Runde {round_num} — Konsens: {consensus:.1%}</h3>"
                for ao in rd.get("agent_outputs", []):
                    role = ao.get("role", "unbekannt")
                    content = esc(ao.get("content", ""))
                    tokens = ao.get("tokens_used", 0)
                    duration_ms = ao.get("duration_ms", 0)
                    llm_pid = ao.get("llm_profile_id", "")
                    display_role = esc(_display_agent_role(role))
                    if llm_pid:
                        display_role += f' <span style="font-weight:normal;color:#666;">— {esc(llm_pid)}</span>'
                    rounds_html += f'<div class="agent-block"><div class="agent-role">{display_role}</div><div class="agent-content">{content}</div>'
                    meta_parts = []
                    if tokens:
                        meta_parts.append(f"Tokens: {tokens}")
                    if duration_ms:
                        meta_parts.append(f"Latenz: {duration_ms}ms")
                    if llm_pid:
                        meta_parts.append(f"LLM-Profil: {esc(llm_pid)}")
                    if meta_parts:
                        rounds_html += f'<div class="agent-meta">{" | ".join(meta_parts)}</div>'
                    rounds_html += "</div>"

        # --- Final output ---
        output = transcript["output"]
        output_section = ""
        if output:
            output_section = f'<h2>Ergebnis</h2><div class="output-text">{esc(output)}</div>'

        # --- Audit trail ---
        audit_rows = ""
        for e in audit_entries:
            audit_rows += (
                f"<tr>"
                f"<td>{esc(str(e.get('timestamp', '')))}</td>"
                f"<td>{esc(str(e.get('event_type', '')))}</td>"
                f"<td>{esc(str(e.get('node_id', '')))}</td>"
                f"<td>{esc(str(e.get('actor', '')))}</td>"
                f"<td>{esc(str(e.get('llm_profile_id', '')))}</td>"
                f"<td>{e.get('latency_ms', 0)}</td>"
                f"<td>{e.get('prompt_tokens', 0) + e.get('completion_tokens', 0)}</td>"
                f"</tr>"
            )
        audit_section = ""
        if audit_entries:
            audit_section = f"""<h2>Audit-Trail</h2>
<table>
<thead><tr><th>Zeitstempel</th><th>Ereignis</th><th>Knoten</th><th>Aktor</th><th>LLM-Profil</th><th>Latenz (ms)</th><th>Tokens</th></tr></thead>
<tbody>{audit_rows}</tbody>
</table>"""

        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>Bericht — {title}</title>
<style>
  body {{ font-family: Calibri, 'Segoe UI', sans-serif; margin: 2em; color: #222; line-height: 1.6; }}
  h1 {{ color: #1a1a2e; border-bottom: 2px solid #16213e; padding-bottom: 0.3em; }}
  h2 {{ color: #16213e; margin-top: 1.5em; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }}
  h3 {{ color: #0f3460; margin-top: 1.2em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 0.5em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; }}
  th {{ background: #f0f0f0; }}
  .case-text {{ background: #f8f9fa; padding: 1em; border-left: 4px solid #0f3460; margin: 0.5em 0; white-space: pre-wrap; }}
  .agent-block {{ margin: 1em 0; padding: 0.8em; border: 1px solid #e0e0e0; border-radius: 6px; background: #fafafa; }}
  .agent-role {{ font-weight: bold; color: #0f3460; margin-bottom: 0.4em; font-size: 1.05em; }}
  .agent-content {{ white-space: pre-wrap; }}
  .agent-meta {{ font-size: 0.85em; color: #888; margin-top: 0.4em; }}
  .output-text {{ background: #e8f5e9; padding: 1em; border-left: 4px solid #2e7d32; white-space: pre-wrap; }}
  .meta-table {{ width: auto; }}
  .meta-table td {{ padding: 4px 12px; }}
</style></head><body>
<h1>{title}</h1>
<table class="meta-table">{meta_rows}</table>
{case_section}
{rounds_html}
{output_section}
{audit_section}
</body></html>"""
