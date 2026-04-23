import asyncio
import html
import json
import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS
from ..core.debate_engine import DebateState

logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self):
        self.output_dir = Path("reports")
        self.output_dir.mkdir(exist_ok=True)

    async def generate(self, state: DebateState, fmt: str = "docx") -> Path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"debate_{state.session_id}_{ts}.{fmt}"
        path = self.output_dir / filename

        if fmt == "docx":
            await asyncio.to_thread(self._build_docx, state, path)
        elif fmt == "pdf":
            await asyncio.to_thread(self._build_pdf, state, path)
        else:
            raise ValueError("Unterstützte Formate: 'docx', 'pdf'")

        logger.info(f"📄 Report generiert: {path}")
        return path

    def _build_docx(self, state: DebateState, path: Path):
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        
        # Titel
        h = doc.add_heading(f"Debatte-Analyse: {state.session_id}", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadaten
        doc.add_heading("Metadaten & Prozess", level=1)
        t = doc.add_table(rows=5, cols=2)
        data = [
            ("Erstellt am", state.created_at),
            ("Durchlaufene Runden", str(len(state.rounds))),
            ("Finale Konsens-Bewertung", f"{state.final_consensus:.2f}"),
            ("Externe Validierung", "Aktiv" if state.validation_report else "Deaktiviert"),
            ("Präzedenzfälle konsultiert", str(len(state.precedents_retrieved)))
        ]
        for i, (k, v) in enumerate(data):
            t.cell(i, 0).text = k
            t.cell(i, 1).text = v

        # Sachverhalt
        doc.add_heading("Sachverhalt / Input", level=1)
        doc.add_paragraph(state.context[:1200] + ("… [gekürzt]" if len(state.context) > 1200 else ""))

        # Finale Argumentation
        doc.add_heading("Finale Argumentation / Strategie", level=1)
        for line in state.output.split('\n'):
            doc.add_paragraph(line.replace('-', '•') if line.strip().startswith('-') else line)

        # Validierung
        if state.validation_report:
            doc.add_heading("Externe Faktenprüfung", level=1)
            for v in state.validation_report:
                p = doc.add_paragraph()
                p.add_run(f"Behauptung: ").bold = True
                p.add_run(v["claim"])
                for ev in v.get("evidence", [])[:2]:
                    doc.add_paragraph(f"  • {ev.get('title', 'Unbekannt')}: {ev.get('snippet', '')[:120]}…", style='List Bullet')

        # Audit-Referenz
        doc.add_heading("Audit & Reproduzierbarkeit", level=1)
        doc.add_paragraph(
            f"Vollständiger Trace: logs/{state.session_id}.jsonl\n"
            f"Prompt-Versionen: Siehe Trace-Header\n"
            f"Generiert durch: Debate-Agent v0.1.0 | {datetime.now().isoformat()}"
        )
        doc.save(path)

    def _build_pdf(self, state: DebateState, path: Path):
        html_safe = lambda s: html.escape(s).replace('\n', '<br>')
        
        validation_html = ""
        if state.validation_report:
            val_items = []
            for v in state.validation_report:
                ev_html = "".join(
                    f"<li><b>{html.escape(e.get('title','?'))}:</b> {html.escape(e.get('snippet','')[:100])}…</li>"
                    for e in v.get("evidence", [])[:2]
                )
                val_items.append(f"<p><b>Claim:</b> {html_safe(v['claim'])}</p><ul>{ev_html}</ul>")
            validation_html = "<h2>Externe Faktenprüfung</h2>" + "".join(val_items)

        content = f"""
        <!DOCTYPE html>
        <html><head><meta charset="utf-8">
        <style>
            body {{ font-family: 'DejaVu Sans', 'Segoe UI', sans-serif; line-height: 1.6; margin: 2.5cm; color: #222; }}
            h1 {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 10px; font-size: 22pt; }}
            h2 {{ color: #1a365d; margin-top: 28px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
            .meta {{ background: #f4f6f8; padding: 12px 15px; border-radius: 4px; margin-bottom: 20px; }}
            .meta table {{ width: 100%; border-collapse: collapse; }}
            .meta td {{ padding: 4px 8px; border-bottom: 1px solid #e0e0e0; }}
            .meta td:first-child {{ font-weight: bold; width: 35%; }}
            .footer {{ margin-top: 60px; font-size: 0.75em; color: #666; text-align: center; border-top: 1px solid #bbb; padding-top: 10px; }}
        </style></head><body>
            <h1>Debatte-Analyse: {state.session_id}</h1>
            <div class="meta"><table>
                <tr><td>Erstellt am</td><td>{state.created_at}</td></tr>
                <tr><td>Runden</td><td>{len(state.rounds)}</td></tr>
                <tr><td>Konsens</td><td>{state.final_consensus:.2f}</td></tr>
                <tr><td>Präzedenzfälle</td><td>{len(state.precedents_retrieved)}</td></tr>
            </table></div>
            <h2>Sachverhalt</h2><p>{html_safe(state.context[:1200])}{"… [gekürzt]" if len(state.context) > 1200 else ""}</p>
            <h2>Finale Argumentation</h2><p>{html_safe(state.output)}</p>
            {validation_html}
            <div class="footer">Audit-Trace: logs/{state.session_id}.jsonl | Generiert: {datetime.now().isoformat()} | Debate-Agent v0.1.0</div>
        </body></html>
        """
        HTML(string=content).write_pdf(str(path))